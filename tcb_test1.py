#pip install openpyxl
#pip install python-docx

import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import openpyxl
import os
from docx import Document
from docx.oxml import OxmlElement

def load_excel_data(file_path):
    wb = openpyxl.load_workbook(file_path)
    sheet = wb.active
    data = {}
    
    headers = [cell.value for cell in sheet[1]]
    for row in sheet.iter_rows(min_row=2, values_only=True):
        key = row[0]
        data[key] = {headers[i]: (row[i] if row[i] is not None else '') for i in range(1, len(row))}
    # for i in range(1,len(row)):
    #     print(row[i])
    return data, headers


def display_excel_data(file_path):
    try:
        data, headers = load_excel_data(file_path)
        text_widget.config(state=tk.NORMAL)
        text_widget.delete('1.0', tk.END)
        
        # 헤더 출력
        header_line = "\t".join(headers)
        text_widget.insert(tk.END, f"{header_line}\n\n")

        # 각 행 데이터 출력
        for key, values in data.items():
            row_data = [key] + [values[header] for header in headers[1:]]
            row_line = "\t".join(map(str, row_data))
            text_widget.insert(tk.END, f"{row_line}\n")
        
        text_widget.config(state=tk.DISABLED)
    except Exception as e:
        messagebox.showerror("오류", f"엑셀 데이터를 불러오는 중 오류가 발생했습니다: {e}")

# def replace_placeholders(word_file, output_file, data, key):
#     try:
#         doc = Document(word_file)
#         for paragraph in doc.paragraphs:
#             for placeholder, value in data[key].items():
#                 print('placeholder :',placeholder,'value:',value )
#                 if placeholder in paragraph.text:
#                     paragraph.text = paragraph.text.replace(placeholder, str(value))
#                     # 글자 대체의 원리라서 {q1}라고 쓰인 부분에 '가' 들어가는 것.
#                     #방법-> 변수명 자체를 {q1} 이런식으로 저장해서, {가}로 대체하기
#         doc.save(output_file)



def replace_placeholders(word_file, output_file, data, key):
    try:
        doc = Document(word_file)

        # 텍스트 대체 함수
        def replace_text(text, placeholders):
            for placeholder, value in placeholders.items():
                if placeholder in text:
                    text = text.replace(placeholder, str(value))
            return text

        # 표 처리 함수
        def process_tables():
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.text = replace_text(paragraph.text, data[key])

        # 일반 텍스트 처리 함수
        def process_paragraphs():
            for paragraph in doc.paragraphs:
                paragraph.text = replace_text(paragraph.text, data[key])

        # 표와 일반 텍스트 모두 처리
        process_tables()
        process_paragraphs()

        # 저장
        doc.save(output_file)

    except Exception as e:
        raise RuntimeError(f"Word 파일을 처리하지 못했습니다: {e}")
    
    
def select_excel_file():
    global excel_file_path
    excel_file_path = filedialog.askopenfilename(filetypes=[("Excel 파일", "*.xlsx")])
    if excel_file_path:
        excel_label.config(text=os.path.basename(excel_file_path))
        display_excel_data(excel_file_path)

def select_word_file():
    global word_file_path
    word_file_path = filedialog.askopenfilename(filetypes=[("Word 파일", "*.docx")])
    if word_file_path:
        word_label.config(text=os.path.basename(word_file_path))

def perform_conversion():
    if not excel_file_path or not word_file_path:
        messagebox.showerror("오류", "Excel 파일과 Word 파일을 모두 선택해 주세요.")
        return

    try:
        excel_data, _ = load_excel_data(excel_file_path)
        key = key_entry.get()
        if key not in excel_data:
            messagebox.showerror("오류", f"Excel 데이터에서 키 '{key}'를 찾을 수 없습니다.")
            return
        output_file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word 파일", "*.docx")])
        replace_placeholders(word_file_path, output_file_path, excel_data, key)
        messagebox.showinfo("성공", f"파일이 {output_file_path}로 저장되었습니다.")
    except Exception as e:
        messagebox.showerror("오류", f"오류가 발생했습니다: {e}")

# GUI 설정
root = tk.Tk()
root.title("Excel to Word 변환기")

excel_file_path = None
word_file_path = None

# 레이블과 버튼 설정
tk.Label(root, text="Excel 파일 선택:").grid(row=0, column=0, padx=10, pady=10, sticky="w")
excel_label = tk.Label(root, text="")
excel_label.grid(row=0, column=2, padx=10, pady=10, sticky="w")
excel_button = tk.Button(root, text="찾아보기", command=select_excel_file)
excel_button.grid(row=0, column=1, padx=10, pady=10, sticky="w")

tk.Label(root, text="Word 파일 선택:").grid(row=1, column=0, padx=10, pady=10, sticky="w")
word_label = tk.Label(root, text="")
word_label.grid(row=1, column=2, padx=10, pady=10, sticky="w")
word_button = tk.Button(root, text="찾아보기", command=select_word_file)
word_button.grid(row=1, column=1, padx=10, pady=10, sticky="w")

tk.Label(root, text="키 입력:").grid(row=2, column=0, padx=10, pady=10, sticky="w")
key_entry = tk.Entry(root)
key_entry.grid(row=2, column=1, padx=10, pady=10, sticky="w")

convert_button = tk.Button(root, text="변환 및 저장", command=perform_conversion)
convert_button.grid(row=3, column=0, columnspan=3, pady=20)

# 텍스트 박스와 스크롤바 설정
text_frame = tk.Frame(root)
text_frame.grid(row=4, column=0, columnspan=3, padx=10, pady=10, sticky="nsew")

x_scrollbar = tk.Scrollbar(text_frame, orient=tk.HORIZONTAL)
x_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)

y_scrollbar = tk.Scrollbar(text_frame, orient=tk.VERTICAL)
y_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

text_widget = tk.Text(text_frame, wrap=tk.NONE, width=80, height=20, xscrollcommand=x_scrollbar.set, yscrollcommand=y_scrollbar.set)
text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

x_scrollbar.config(command=text_widget.xview)
y_scrollbar.config(command=text_widget.yview)

# 창 크기 조정 시 위젯 크기 조정
root.grid_rowconfigure(4, weight=1)
root.grid_columnconfigure(0, weight=1)
root.grid_columnconfigure(1, weight=0)  # 변하지 않는 고정 크기
root.grid_columnconfigure(2, weight=1)

root.mainloop()
